import os
import uuid
import hashlib
from dotenv import load_dotenv
import google.generativeai as genai
from fastapi import FastAPI, Form, UploadFile, File, HTTPException, Depends
from fastapi.responses import JSONResponse
from langchain_core.prompts import PromptTemplate
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pydantic import BaseModel
import torch
from docx import Document
from fastapi_cache import FastAPICache
from fastapi_cache.backends.inmemory import InMemoryBackend
from fastapi_cache.decorator import cache
import shutil

app = FastAPI()

# Load environment variables from .env file
load_dotenv()

# Set up the API key for Gemini
api_key = os.getenv('GOOGLE_API_KEY')
if not api_key:
    raise ValueError("No API key found. Please set the GOOGLE_API_KEY environment variable.")

# Configure the Gemini model
genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-1.5-pro')  # Ensure this model name is correct

# Configuration variables
DATA_PATH = os.getenv('DATA_PATH', 'docs/')
DB_FAISS_PATH = os.getenv('DB_FAISS_PATH', 'vectorstore_doc/db_faiss')
HASH_FILE_PATH = os.getenv('HASH_FILE_PATH', 'vectorstore_doc/document_hashes.txt')
ALLOWED_EXTENSIONS = {'.pdf', '.txt', '.docx'}

CUSTOM_PROMPT_TEMPLATE = """Use the following pieces of information to answer the user's question or provide a summary of the document.
If you don't know the answer or summary, just say that you don't know, don't try to make up an answer.
Context: {context}
Question or request for document summary: {question}
If the question is asking for a summary, provide a concise overview of the main points in the document.
If it's a specific question, answer it based on the information in the context.
If asked to compare or analyze, use the relevant information from the context to do so.
Only return the helpful answer or document summary below and nothing else.

Helpful answer or document summary:
"""

device = torch.device('mps' if torch.backends.mps.is_available() else 'cpu')

class Query(BaseModel):
    query: str
    history: list = []  # Add history to manage conversational context

# Use UUID for session management
sessions = {}

def get_session(session_id: str):
    if session_id not in sessions:
        user_db_path = f"vectorstore_doc/db_faiss_{session_id}"
        sessions[session_id] = {"history": [], "last_query": None, "db_path": user_db_path}
    return sessions[session_id]

def generate_session_id():
    return str(uuid.uuid4())

def get_file_hash(file_path):
    """Compute a hash for the file to check for duplicates."""
    hash_algo = hashlib.md5()
    with open(file_path, 'rb') as file:
        while chunk := file.read(8192):
            hash_algo.update(chunk)
    return hash_algo.hexdigest()

def load_existing_hashes():
    """Load existing document hashes from a file."""
    if os.path.exists(HASH_FILE_PATH):
        with open(HASH_FILE_PATH, 'r') as file:
            return set(line.strip() for line in file)
    return set()

def save_hashes(hashes):
    """Save updated document hashes to a file."""
    with open(HASH_FILE_PATH, 'w') as file:
        file.write('\n'.join(hashes))

def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return [Document(page_content=text, metadata={"source": file_path})]

def load_document(file_path):
    if file_path.lower().endswith('.pdf'):
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif file_path.lower().endswith('.txt'):
        loader = TextLoader(file_path)
        return loader.load()
    elif file_path.lower().endswith('.docx'):
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

def create_or_update_vector_db(file_path, db_path):
    document_texts = load_document(file_path)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    split_texts = text_splitter.split_documents(document_texts)
    embeddings = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2', model_kwargs={'device': device})

    # Remove the old database if it exists
    if os.path.exists(db_path):
        if os.path.isfile(db_path):
            os.remove(db_path)
        elif os.path.isdir(db_path):
            shutil.rmtree(db_path)

    # Create a new database
    db = FAISS.from_documents(split_texts, embeddings)
    db.save_local(db_path)

    return "Vector database created successfully."

def set_custom_prompt():
    prompt = PromptTemplate(template=CUSTOM_PROMPT_TEMPLATE, input_variables=['context', 'question'])
    return prompt

def generate_response(prompt, include_invitation=False):
    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()

        if include_invitation:
            invitation_prompt = f"""
            Based on the following response:

            {response_text}

            Generate a natural, engaging invitation for the user to continue the conversation.
            This should flow from the topic and encourage further discussion, but phrase it as
            a suggestion or option for the user, starting with phrases like "Would you like to..."
            or "If you're interested, we could...". Make it sound friendly and optional.
            """
            invitation = model.generate_content(invitation_prompt).text.strip()
            return f"{response_text}\n\n{invitation}"
        else:
            return response_text
    except Exception as e:
        print(f"Error generating response: {e}")
        return "I'm sorry, I couldn't generate a response."


def generate_follow_up_question(user_input, assistant_answer, history):
    context = "\n".join(history[-10:])  # Use last 10 exchanges for context
    follow_up_prompt = f"""
    Based on the following conversation:

    {context}

    User's last question: {user_input}
    Assistant's last answer: {assistant_answer}

    Suggest a relevant, engaging follow-up question that the user might ask to continue the conversation.
    """
    return generate_response(follow_up_prompt)

def query_gemini(prompt, conversation_history):
    full_prompt = f"""You are an AI assistant. Here is the ongoing conversation and context:

{conversation_history}

User's current question or request: {prompt}

Please provide a comprehensive and detailed answer based solely on the provided document context. Ensure your response addresses the user's query using the document content, and do not include information outside of the document.

Answer:"""
    return generate_response(full_prompt)

def qa_bot(query, history, db_path):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", model_kwargs={'device': device})

    if os.path.exists(db_path):
        db = FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)
    else:
        return "No document has been uploaded yet. Please upload a document first.", history

    conversation_history = '\n'.join(history[-5:])
    document_context = ' '.join([doc.page_content for doc in db.similarity_search(query, k=2)])
    combined_context = f"Document Context:\n{document_context}\n\nRecent Conversation:\n{conversation_history}\n\nCurrent Query: {query}"

    assistant_answer = generate_response(combined_context, include_invitation=True)

    history.append(f"User: {query}")
    history.append(f"Assistant: {assistant_answer}")

    return assistant_answer, history


@app.on_event("startup")
async def startup():
    FastAPICache.init(InMemoryBackend())


@app.post("/chatwithdoc")
@cache(expire=3600)  # Cache for 1 hour
async def chat_with_doc(file: UploadFile = File(...), session_id: str = Form(...)):
    if not file:
        return JSONResponse({"error": "Empty File provided"}, status_code=400)
    if not any(file.filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
        return JSONResponse({"error": "Unsupported file type"}, status_code=400)

    file_path = os.path.join(DATA_PATH, file.filename)
    with open(file_path, "wb") as buffer:
        buffer.write(file.file.read())

    try:
        session = get_session(session_id)

        # Clear the history
        session["history"] = []

        # Create or update the vector database with the new file
        create_or_update_vector_db(file_path, session["db_path"])

        return JSONResponse({"message": "Document uploaded and vector database updated successfully."})
    except Exception as e:
        return JSONResponse({"error": f"An error occurred: {str(e)}"}, status_code=500)
    finally:
        os.remove(file_path)  # Remove the temporary file after processing



@app.post("/ask")
async def ask(query: str = Form(...), session_id: str = Form(...)):
    try:
        session = get_session(session_id)
        response, updated_history = qa_bot(query, session["history"], session["db_path"])
        follow_up_question = generate_follow_up_question(query, response, updated_history)
        session["history"] = updated_history
        return JSONResponse({
            "response": response,
            "follow_up_question": follow_up_question,
            "history": updated_history
        })
    except Exception as e:
        return JSONResponse({"error": f"An error occurred: {str(e)}"}, status_code=500)

@app.post("/generate_session")
async def generate_new_session():
    try:
        new_session_id = generate_session_id()
        return JSONResponse({"session_id": new_session_id})
    except Exception as e:
        return JSONResponse({"error": f"An error occurred: {str(e)}"}, status_code=500)

if _name_ == "_main_":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
